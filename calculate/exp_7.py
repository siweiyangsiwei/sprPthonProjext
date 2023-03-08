from PyQt5.QtWidgets import QTableWidgetItem
from docx import Document

def data_processing_7_data_calculate_click(self):
    date = self.data_processing_7_date.text()
    num_of_shai_ban_str = self.data_processing_7_data_num.text()
    num_of_pic_str = self.data_processing_7_data_num_of_pic.text()
    height_of_tian_liao_str = self.data_processing_7_data_height.text()
    num_of_shai_ban = 0.0
    height_of_tian_liao = 0.0
    num_of_pic = 0.0
    tower_type = self.data_processing_7_data_type.isChecked()
    # 分别用于填写三个实验的数据
    data1_str = []
    data2_str = []
    data3_str = []
    data1 = []
    data2 = []
    data3 = []
    item = QTableWidgetItem()
    table = None
    R1 = 0.0
    R2 = 0.0
    R3 = 0.0
    q1 = 0.0
    q2 = 0.0
    q3 = 0.0
    # 没有选择,默认筛板塔
    if not tower_type:
        if num_of_shai_ban_str == '' or num_of_pic_str == '':
            print("请填写完信息")
            return
        else:
            # 开始进行筛板塔计算
            table = self.data_processing_7_table_1
            num_of_shai_ban = float(num_of_shai_ban_str)
            num_of_pic = float(self.data_processing_7_data_num_of_pic.text())
            for i in range(table.rowCount()):
                if table.item(i, 0) is None or table.item(i, 1) is None or table.item(i, 2) is None:
                    print('请先正确填写信息,第' + str(i + 1) + '没有填写')
                    return
                else:
                    t1 = table.item(i, 0).data(0)
                    t2 = table.item(i, 1).data(0)
                    t3 = table.item(i, 2).data(0)
                    data1_str.append(t1)
                    data1.append(float(t1))
                    data2_str.append(t2)
                    data2.append(float(t2))
                    data3_str.append(t3)
                    data3.append(float(t3))
            # 计算全塔效率
            ET = num_of_pic - 1 / num_of_shai_ban
            # 回流比
            R1 = data1[10] / data1[9]
            R2 = data2[10] / data2[9]
            R3 = data3[10] / data3[9]
            # 计算 q
            q1 = 1 + (2400 * (90 - data1[6])) / 200
            q2 = 1 + (2400 * (90 - data2[6])) / 200
            q3 = 1 + (2400 * (90 - data3[6])) / 200
            item.setText(str(ET))
    # 选择了,为填料塔
    else:
        if height_of_tian_liao_str == '' or num_of_pic_str == '':
            print("请填写完信息")
            return
        else:
            # 开始进行填料塔计算
            table = self.data_processing_7_table_2
            height_of_tian_liao = float(height_of_tian_liao_str)
            num_of_pic = float(self.data_processing_7_data_num_of_pic.text())
            for i in range(table.rowCount()):
                if table.item(i, 0) is None or table.item(i, 1) is None or table.item(i, 2) is None:
                    print('请先正确填写信息,第' + str(i + 1) + '没有填写')
                    return
                else:
                    t1 = table.item(i, 0).data(0)
                    t2 = table.item(i, 1).data(0)
                    t3 = table.item(i, 2).data(0)
                    data1_str.append(t1)
                    data1.append(float(t1))
                    data2_str.append(t2)
                    data2.append(float(t2))
                    data3_str.append(t3)
                    data3.append(float(t3))
            # 计算等板高度
            HETP = height_of_tian_liao / num_of_pic
            # 回流比
            R1 = data1[13] / data1[12]
            R2 = data2[13] / data2[12]
            R3 = data3[13] / data3[12]
            # 计算 q
            q1 = 1 + (2400 * (90 - (data1[6] + data1[5]) / 2)) / 200
            q2 = 1 + (2400 * (90 - (data2[6] + data3[5]) / 2)) / 200
            q3 = 1 + (2400 * (90 - (data2[6] + data3[5]) / 2)) / 200
            item.setText(str(HETP))
    table3 = self.data_processing_7_table_3
    table3.setItem(0, 0, item)
    item = QTableWidgetItem()
    item.setText(str(R1))
    table3.setItem(1, 0, item)
    item = QTableWidgetItem()
    item.setText(str(R2))
    table3.setItem(1, 1, item)
    item = QTableWidgetItem()
    item.setText(str(R3))
    table3.setItem(1, 2, item)
    item = QTableWidgetItem()
    item.setText(str(q1))
    table3.setItem(2, 0, item)
    item = QTableWidgetItem()
    item.setText(str(q2))
    table3.setItem(2, 1, item)
    item = QTableWidgetItem()
    item.setText(str(q3))
    table3.setItem(2, 2, item)

    # 实验报告
    document = Document('./resources/report/实验七 精馏实验.docx')
    document.add_paragraph("                表1 数据记录表")
    table1 = document.add_table(len(data1) + 1, 4, style='Table Grid')
    table1.rows[0].cells[1].text = '1'
    table1.rows[0].cells[2].text = '2'
    table1.rows[0].cells[3].text = '3'
    if not tower_type:
        for i in range(len(data1)):
            table1.rows[i + 1].cells[0].text = self.data_processing_7_table_1.verticalHeaderItem(i).data(0)

    else:
        for i in range(len(data1)):
            table1.rows[i + 1].cells[0].text = self.data_processing_7_table_2.verticalHeaderItem(i).data(0)

    for j in range(len(data1)):
        table1.rows[j + 1].cells[1].text = str(data1[j])
        table1.rows[j + 1].cells[2].text = str(data2[j])
        table1.rows[j + 1].cells[3].text = str(data3[j])
    document.add_paragraph("\n")

    document.add_paragraph("                表2 计算结果")
    table2 = document.add_table(4, 4, style='Table Grid')
    table2.rows[0].cells[1].text = '1'
    table2.rows[0].cells[2].text = '2'
    table2.rows[0].cells[3].text = '3'
    for i in range(3):
        table2.rows[i + 1].cells[0].text = self.data_processing_7_table_3.verticalHeaderItem(i).data(0)
        table2.rows[1].cells[1].text = str(HETP)
        table2.rows[1].cells[2].text = str(HETP)
        table2.rows[1].cells[3].text = str(HETP)
        table2.rows[2].cells[1].text = str(R1)
        table2.rows[2].cells[2].text = str(R2)
        table2.rows[2].cells[3].text = str(R3)
        table2.rows[3].cells[1].text = str(q1)
        table2.rows[3].cells[2].text = str(q2)
        table2.rows[3].cells[3].text = str(q3)
    document.save('实验七 精馏实验.docx')
