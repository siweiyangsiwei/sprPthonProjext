import math
import os

from PyQt5.QtWidgets import QTableWidgetItem
from PyQt5 import QtGui
import matplotlib.pyplot as plt
from docx import Document


def data_processing_8_data_calculate_click(self):
    if (self.data_processing_8_date.text() == ''
            or self.data_processing_8_data_height.text() == ''
            or self.data_processing_8_data_diameter.text() == ''
            or self.data_processing_8_data_specification.text() == ''
            or self.data_processing_8_data_pressure.text() == ''
            or self.data_processing_8_data_temp.text() == ''
            or self.data_processing_8_data_T.text() == ''
            or self.data_processing_8_data_P.text() == ''):
        print("请填写完整的信息")
        return
    # 实验日期
    date = self.data_processing_8_date.text()
    # 填料层高度
    height_str = self.data_processing_8_data_height.text()
    height = float(self.data_processing_8_data_height.text())
    # 塔内径
    diameter_str = self.data_processing_8_data_diameter.text()
    diameter = float(self.data_processing_8_data_diameter.text())
    # 填料规格
    specification = self.data_processing_8_data_specification.text()
    # 大气压
    pressure_str = self.data_processing_8_data_pressure.text()
    pressure = float(self.data_processing_8_data_pressure.text())
    # 室温
    temp_str = self.data_processing_8_data_temp.text()
    temp = float(self.data_processing_8_data_temp.text())
    # 流量计标定状态
    T = self.data_processing_8_data_T.text()
    P = self.data_processing_8_data_P.text()
    # 记录填料塔流体力学性能
    data = []
    table = None
    # 没有选择,填料为干料
    if not self.data_processing_8_data_type.isChecked():
        table = self.data_processing_8_table_1
        for i in range(table.rowCount()):
            data.append([])
            for j in range(table.columnCount()):
                if table.item(i, j) is None:
                    if j == 1:
                        item = QTableWidgetItem()
                        item.setText(str(data[i][0] / height))
                        table.setItem(i, j, item)
                        data[i].append(float(table.item(i, j).data(0)))
                    else:
                        print("请先填写好表中的数据")
                        return
                else:
                    data[i].append(float(table.item(i, j).data(0)))
    # 填料为湿料
    else:
        table = self.data_processing_8_table_2
        for i in range(table.rowCount()):
            for j in range(table.columnCount() - 1):
                if table.item(i, j) is None:
                    if j == 1:
                        item = QTableWidgetItem()
                        item.setText(str(data[i][0] / height))
                        table.setItem(i, j, item)
                        data[i].append(float(table.item(i, j).data(0)))
                    else:
                        print("请先填写好表中的数据")
                        return
                else:
                    data[i].append(float(table.item(i, j).data(0)))
    # 获取到数据了,开始进行塔的流体力学性能计算
    # 计算V0标准状态下空气流量的计算 以及 空塔气速的计算
    V0 = []
    u = []
    # 获得一个压强降数组
    deter_pressure = []
    for i in range(table.rowCount()):
        V = data[i][2] * 273.15 / 101.32 * math.sqrt(101.32 * data[i][3] / (273.15 + 20) / (273.15 + data[i][4]))
        u1 = 4 * V / math.pi / height.__pow__(2)
        V0.append(V)
        u.append(u1)
        deter_pressure.append(data[i][1])

    # 绘制不同喷淋量下填料层的压强降与气速的关系
    plt.plot(u, deter_pressure, color='r', marker='o')
    plt.xlabel('u')
    plt.ylabel('△p')
    plt.title('△P/Z-u curve')
    plt.savefig('./data/img/exp_8_data_1.png')
    img = os.path.abspath('./data/img/exp_8_data_1.png')
    image = QtGui.QPixmap(img).scaled(11182, 1182)
    self.data_processing_8_data_pic_1.setScaledContents(True)
    self.data_processing_8_data_pic_1.setPixmap(image)

    table = self.data_processing_8_table_3
    # 用于保存表3的数据
    data3 = [[], []]
    # 进行CO2传质性能测试数据计算
    for i in range(table.rowCount()):
        for j in range(2):
            if table.item(i, j) is None:
                print("请填写好数据再进行计算")
                return
            else:
                data3[j].append(float(table.item(i, j).data(0)))

    # 对于表2进行表3的数据计算
    # 用于保存表4的计算数据(浮点数类型)
    data4 = [[], []]
    for i in range(2):
        data4[i].append(data3[i][0])
        # CO2体积流量
        Vco2 = data3[i][1] * data3[i][7]
        # 计算标准状态下CO2的体积流量
        Vco20 = Vco2 * 273.15 / 101.32 * math.sqrt(
            1.293 / 1.977 * 101.32 * data3[i][3] / (273.15 + 20) / (data3[i][2] + 273.15))
        data4[i].append(Vco20)
        # 空气体积流量
        Vspace = data3[i][1] * (1 - data3[i][7])
        # 标准状态下空气的体积流量
        Vspace0 = Vspace * 273.15 / 101.32 * math.sqrt(
            101.32 * data3[i][3] / (273.15 + 20) / (273.15 + data3[i][2]))
        data4[i].append(Vspace0)
        # 计算空气摩尔流量
        Vmol = Vspace / 22.4 / 1000 / diameter.__pow__(2)
        data4[i].append(Vmol)
        # 计算水摩尔流量
        Lmol = data3[i][0] / 1000 * 1 / diameter.__pow__(2)
        data4[i].append(Lmol)
        # 计算相平衡常数
        m = data3[i][6] / data3[i][3]
        data4[i].append(m)
        # 计算吸收因数
        A = Lmol / m / Vmol
        # 计算X2
        X2 = Vmol * (data3[i][7] - data3[i][8]) / Lmol
        # 计算N(OL)
        NOL = 1 / (1 - A) * math.log(((1 - A) * (data3[i][7] - m * X2 / data3[i][7])), math.e)
        data4[i].append(NOL)
        # 计算H(OL)
        HOL = height / NOL
        data4[i].append(HOL)
        # 计算Kxa
        Kxa = Lmol / HOL / (0.25 * diameter.__pow__(2) * math.pi)
        data4[i].append(Kxa)
        # 计算吸收率
        rate = (data3[i][7] - data3[i][8]) / data3[i][7]
        data4[i].append(rate)
    # 转换为字符串后填入到表4中
    table = self.data_processing_8_table_4
    for i in range(table.rowCount()):
        for j in range(table.columnCount()):
            item = QTableWidgetItem()
            item.setText(str(data4[j][i]))
            table.setItem(i, j, item)

    # 生成实验报告
    document = Document('./resources/report/实验八 吸收-解吸实验.docx')

    document.add_paragraph("                表1 填料塔流体力学性能")
    table1 = document.add_table(len(data) + 1, len(data[0]), style='Table Grid')
    heading_cells = table1.rows[0].cells
    heading_cells[0].text = "填料层压降KPa"
    heading_cells[1].text = "单位高度填料层压降Kpa"
    heading_cells[2].text = "空气转子流量计读数m"
    heading_cells[3].text = "空气转子流量计前压表KPa"
    heading_cells[4].text = "空气转子流量计前温度℃"
    heading_cells[5].text = "空塔气速m/s"
    # if len(data) > 7:
    #     heading_cells[6].text = "操作现象"

    document.add_paragraph("\n")
    for i in range(len(data)):
        for j in range(len(data[0])):
            table1.rows[i + 1].cells[j].text = str(data[i][j])

    document.add_paragraph("                表2 CO₂吸收传质系数测定原始数据记录表")
    table3 = document.add_table(len(data3[0]), 3, style='Table Grid')
    table3.rows[0].cells[0].text = '水体积流量L/(L/h)'
    table3.rows[1].cells[0].text = '空气转子流量计读数V₁/(m³/h)'
    table3.rows[2].cells[0].text = '进口处空气温度/℃'
    table3.rows[3].cells[0].text = '空气转子流量计表压/KPa'
    table3.rows[4].cells[0].text = '水进口温度/℃'
    table3.rows[5].cells[0].text = '水出口温度/℃'
    table3.rows[6].cells[0].text = '亨利常数E×10e-5'
    table3.rows[7].cells[0].text = '进口气体中CO₂浓度Y₁'
    table3.rows[8].cells[0].text = '尾气中CO₂浓度Y₂'
    for i in range(len(data3)):
        for j in range(len(data3[0])):
            table3.rows[j].cells[1 + i].text = str(data3[i][j])
    document.add_paragraph("\n")

    document.add_paragraph("                表3 CO₂吸收传质系数测定数据处理表")
    table4 = document.add_table(len(data4[0]), 3, style='Table Grid')
    table4.rows[0].cells[0].text = '水体积流量L/(L/h)'
    table4.rows[1].cells[0].text = '标准状态下CO₂体积流量V(CO₂)/(m³/h)'
    table4.rows[2].cells[0].text = '标准状态下空气体积流量/(m³/h)'
    table4.rows[3].cells[0].text = '空气摩尔流量V/(kmol/m²·h)'
    table4.rows[4].cells[0].text = '水流量L/(kmol/m²/h)'
    table4.rows[5].cells[0].text = '相平衡常数m'
    table4.rows[6].cells[0].text = 'N(OL)'
    table4.rows[7].cells[0].text = 'H(OL)'
    table4.rows[8].cells[0].text = '液相体积传质系数K(X)α[kmol/(m³·h)]'
    table4.rows[9].cells[0].text = '吸收率%'
    for i in range(len(data4)):
        for j in range(len(data4[0])):
            table4.rows[j].cells[1 + i].text = str(data4[i][j])
    document.add_paragraph('\n')
    document.add_picture('./data/img/exp_8_data_1.png')

    document.save('./resources/report/实验八 吸收-解吸实验.docx')
