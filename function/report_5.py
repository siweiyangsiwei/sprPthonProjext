from PyQt5.QtWidgets import QFileDialog, QMessageBox
from win32com import client
from docx import Document


def import_data(self):
    file_path = QFileDialog.getExistingDirectory()
    return file_path

def write_5_docx(self, p1, p2, p3, t1, t2, t3, g1, g2, g3, K1, K2, K3, qe1, qe2, qe3, te1, te2, te3, qt1, qt2, qt3, s,
                 r, V, u, date, area, temp, q1, q2, q3):
    # 实验报告
    document = Document('./resources/report/实验五 恒压过滤实验.docx')
    document.add_paragraph('实验日期' + str(date) + "\t室温" + str(temp) + "\t过滤面积" + str(area) +
                           "\t滤渣压缩系数" + str(s) + "\t滤液的粘度" + str(u) +
                           "\t滤渣比阻" + str(r) + "\t单位滤液体积的滤渣体积" + str(V))
    document.add_paragraph("                表1 数据记录表")
    table1 = document.add_table(30 + 1, 3, style='Table Grid')

    # 设置表格样式
    for i in range(3):
        table1.rows[0].cells[i].text = self.data_processing_5_table_1.horizontalHeaderItem(i).data(0)

    table1.cell(1, 0).merge(table1.cell(10, 0))
    table1.cell(11, 0).merge(table1.cell(20, 0))
    table1.cell(21, 0).merge(table1.cell(30, 0))

    # 填入数据

    table1.cell(1, 0).text = str(p1)
    table1.cell(11, 0).text = str(p2)
    table1.cell(21, 0).text = str(p3)

    for j in range(len(t1)):
        table1.cell(j + 1, 1).text = str(t1[j])
        table1.cell(11 + j, 1).text = str(t2[j])
        table1.cell(21 + j, 1).text = str(t3[j])

        table1.cell(j + 1, 2).text = str(g1[j])
        table1.cell(11 + j, 2).text = str(g2[j])
        table1.cell(21 + j, 2).text = str(g3[j])

    document.add_paragraph("\n")

    document.add_paragraph("                表2 数据处理表")
    table2 = document.add_table(31, 7, style='Table Grid')
    for i in range(7):
        table2.rows[0].cells[i].text = self.data_processing_5_table_2.horizontalHeaderItem(i).data(0)

    for i in [0, 4, 5, 6]:
        table2.cell(1, i).merge(table2.cell(10, i))
        table2.cell(11, i).merge(table2.cell(20, i))
        table2.cell(21, i).merge(table2.cell(30, i))

    table2.cell(1, 0).text = str(p1)
    table2.cell(11, 0).text = str(p2)
    table2.cell(21, 0).text = str(p3)
    table2.cell(1, 4).text = str(K1)
    table2.cell(11, 4).text = str(K2)
    table2.cell(21, 4).text = str(K3)
    table2.cell(1, 5).text = str(qe1)
    table2.cell(11, 5).text = str(qe2)
    table2.cell(21, 5).text = str(qe3)
    table2.cell(1, 6).text = str(te1)
    table2.cell(11, 6).text = str(te2)
    table2.cell(21, 6).text = str(te3)
    for i in range(10):
        table2.cell(i + 1, 1).text = str(t1[i])
        table2.cell(11 + i, 1).text = str(t2[i])
        table2.cell(21 + i, 1).text = str(t3[i])
        table2.cell(i + 1, 2).text = str(q1[i])
        table2.cell(11 + i, 2).text = str(q2[i])
        table2.cell(21 + i, 2).text = str(q3[i])
        table2.cell(i + 1, 3).text = str(qt1[i])
        table2.cell(11 + i, 3).text = str(qt2[i])
        table2.cell(21 + i, 3).text = str(qt3[i])
    document.add_paragraph("\n")
    document.add_picture('./data/img/exp_5_data_1.png')



    # 保存文件
    pre_path = import_data(self)
    file_path = pre_path + '/' + '实验五 恒压过滤实验.docx'
    document.save(file_path)
    QMessageBox.information(self, '提示', '导出成功', QMessageBox.Ok)
