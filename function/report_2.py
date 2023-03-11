import docx
from PyQt5.QtWidgets import QMessageBox,QFileDialog
from win32com import client

# 实验二 导出实验报告

# 获取实验数据
def get_data(self):
    write_docx(self)

# 选择文件要保存的位置
def import_data(self):
    file_path = QFileDialog.getExistingDirectory()
    return file_path

# 生成docx文件
def write_docx(self):
    try:
        pre_path = import_data(self)
        if pre_path == '':
            pass
        else:
            # 获取模板
            doc = docx.Document('resources/report/实验二 机械能转化实验.docx')

            # 写入实验数据
            doc.add_paragraph("管内径："+ self.d2.text() + "mm" + '\t' + "变截面管喉径：" + self.d_2.text() + 'mm' + '\t' + "水温" + self.t2.text() + "℃" + '\t' + "密度：" + self.rou2.text() + 'kg／m3' + '\t' + "粘度：" + self.miu2.text() + 'Pa·s')

            # 表1
            doc.add_paragraph("")
            doc.add_paragraph("表1 机械能转化实验原始数据记录表")
            # 获取行列
            table_row = self.exp_data_2a.rowCount()
            table_column = self.exp_data_2a.columnCount()
            word_row = self.exp_data_2a.rowCount()+1
            word_column = self.exp_data_2a.columnCount()+1

            # 设置表格 Table Grid网格；Light Shading三线表
            table = doc.add_table(word_row,word_column,style="Table Grid")
            # 写入表头
            table.cell(0,0).text = "序号"
            # 行
            for i in range(table_row):
                # 写入表头数据
                table.cell(i+1, 0).text = str(i + 1)
            # 列
            for i in range(table_column):
                table.cell(0, i+1).text = self.exp_data_2a.horizontalHeaderItem(i).text()
            # 写入内容
            for i in range(table_row):
                for j in range(table_column):
                    if(self.exp_data_2a.item(i,j) != None):
                        table.cell(i+1,j+1).text = self.exp_data_2a.item(i,j).text()

            # 表2
            doc.add_paragraph("")
            doc.add_paragraph("表2 实验数据处理")
            # 获取行列
            table_row_2 = self.exp_data_2b.rowCount()
            table_column_2 = self.exp_data_2b.columnCount()
            word_row_2 = self.exp_data_2b.rowCount() + 1
            word_column_2 = self.exp_data_2b.columnCount() + 1
            table_2 = doc.add_table(word_row_2, word_column_2, style="Table Grid")
            # 写入表头
            table_2.cell(0, 0).text = "序号"
            # 行
            for i in range(table_row_2):
                # 写入表头数据
                table_2.cell(i + 1, 0).text = str(i + 1)
            # 列
            for i in range(table_column_2):
                table_2.cell(0, i + 1).text = self.exp_data_2b.horizontalHeaderItem(i).text()
            # 写入内容
            for i in range(table_row_2):
                for j in range(table_column_2):
                    if (self.exp_data_2b.item(i, j) != None):
                        table_2.cell(i + 1, j + 1).text = self.exp_data_2b.item(i, j).text()

            # 获取保存路径
            file_path = pre_path + '/' + '实验二 机械能转化实验.docx'
            doc.save(file_path)
            QMessageBox.information(self,'提示','导出成功',QMessageBox.Ok)

    except:
        QMessageBox.critical(self, '错误', '导出失败', QMessageBox.Ok)
