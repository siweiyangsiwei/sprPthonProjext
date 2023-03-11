from PyQt5.QtWidgets import QFileDialog, QMessageBox
from docx import Document


def import_data(self):
    file_path = QFileDialog.getExistingDirectory()
    return file_path

def write_8_docx(self, data, data3, data4, date, height, diameter, pressure, temp, T, P):
    # 生成实验报告
    document = Document('./resources/report/实验八 吸收-解吸实验.docx')

    document.add_paragraph('实验日期' + str(date) + "\t填料层高度" + str(height) +
                           "\t塔内径" + str(diameter) +
                           "\t大气压" + str(pressure) + "\t室温" + str(temp) +
                           "\t流量计标定状态:T=" + str(T) + " P=" + str(P))

    document.add_paragraph("                表1 填料塔流体力学性能")
    table1 = document.add_table(len(data) + 1, len(data[0]), style='Table Grid')
    heading_cells = table1.rows[0].cells
    heading_cells[0].text = "填料层压降KPa"
    heading_cells[1].text = "单位高度填料层压降Kpa"
    heading_cells[2].text = "空气转子流量计读数m"
    heading_cells[3].text = "空气转子流量计前压表KPa"
    heading_cells[4].text = "空气转子流量计前温度℃"
    heading_cells[5].text = "空塔气速m/s"
    # if len(data) > 7:
    #     heading_cells[6].text = "操作现象"

    document.add_paragraph("\n")
    for i in range(len(data)):
        for j in range(len(data[0])):
            table1.rows[i + 1].cells[j].text = str(data[i][j])

    document.add_paragraph("                表2 CO₂吸收传质系数测定原始数据记录表")
    table3 = document.add_table(len(data3[0]), 3, style='Table Grid')
    table3.rows[0].cells[0].text = '水体积流量L/(L/h)'
    table3.rows[1].cells[0].text = '空气转子流量计读数V₁/(m³/h)'
    table3.rows[2].cells[0].text = '进口处空气温度/℃'
    table3.rows[3].cells[0].text = '空气转子流量计表压/KPa'
    table3.rows[4].cells[0].text = '水进口温度/℃'
    table3.rows[5].cells[0].text = '水出口温度/℃'
    table3.rows[6].cells[0].text = '亨利常数E×10e-5'
    table3.rows[7].cells[0].text = '进口气体中CO₂浓度Y₁'
    table3.rows[8].cells[0].text = '尾气中CO₂浓度Y₂'
    for i in range(len(data3)):
        for j in range(len(data3[0])):
            table3.rows[j].cells[1 + i].text = str(data3[i][j])
    document.add_paragraph("\n")

    document.add_paragraph("                表3 CO₂吸收传质系数测定数据处理表")
    table4 = document.add_table(len(data4[0]), 3, style='Table Grid')
    table4.rows[0].cells[0].text = '水体积流量L/(L/h)'
    table4.rows[1].cells[0].text = '标准状态下CO₂体积流量V(CO₂)/(m³/h)'
    table4.rows[2].cells[0].text = '标准状态下空气体积流量/(m³/h)'
    table4.rows[3].cells[0].text = '空气摩尔流量V/(kmol/m²·h)'
    table4.rows[4].cells[0].text = '水流量L/(kmol/m²/h)'
    table4.rows[5].cells[0].text = '相平衡常数m'
    table4.rows[6].cells[0].text = 'N(OL)'
    table4.rows[7].cells[0].text = 'H(OL)'
    table4.rows[8].cells[0].text = '液相体积传质系数K(X)α[kmol/(m³·h)]'
    table4.rows[9].cells[0].text = '吸收率%'
    for i in range(len(data4)):
        for j in range(len(data4[0])):
            table4.rows[j].cells[1 + i].text = str(data4[i][j])
    document.add_paragraph('\n')
    document.add_picture('./data/img/exp_8_data_1.png')

    # 保存文件
    pre_path = import_data(self)
    file_path = pre_path + '/' + '实验八 吸收-解吸实验.docx'
    document.save(file_path)
    QMessageBox.information(self, '提示', '导出成功', QMessageBox.Ok)
