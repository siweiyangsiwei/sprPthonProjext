import docx
from PyQt5.QtWidgets import QMessageBox,QFileDialog
from win32com import client

# 实验二 导出实验报告

# 获取实验数据
def get_data(self):
    write_docx(self)

# 选择文件要保存的位置
def import_data(self):
    file_path = QFileDialog.getExistingDirectory()
    return file_path

# 转换docx为pdf
def docx2pdf(fn):
    word = client.Dispatch("Word.Application")  # 打开word应用程序
    # for file in files:
    doc = word.Documents.Open(fn)  # 打开word文件
    doc.SaveAs("{}.pdf".format(fn[:-5]), 17)  # 另存为后缀为".pdf"的文件，其中参数17表示为pdf
    doc.Close()  # 关闭原来word文件
    word.Quit()

# 生成docx文件
def write_docx(self):
    # try:
        pre_path = import_data(self)
        if pre_path == '':
            pass
        else:
            # 获取模板
            doc = docx.Document('resources/report/实验三 流体阻力的测定.docx')

            # 写入实验数据
            doc.add_paragraph("管径："+ self.d3.text() + "mm" + '\t' + "管长：" + self.l3.text() + 'm' + '\t' + "温度" + self.t3.text() + "℃" + '\t' + "流体密度：" + self.rou3.text() + 'kg／m3' + '\t' + "流体粘度：" + self.miu3.text() + 'Pa·s')

            # 表1
            doc.add_paragraph("")
            doc.add_paragraph("表1 光滑管数据表")
            # 获取行列
            table_row = self.exp_data_3a.rowCount()
            table_column = self.exp_data_3a.columnCount()
            word_row = self.exp_data_3a.rowCount()+1
            word_column = self.exp_data_3a.columnCount()+1

            # 设置表格 Table Grid网格；Light Shading三线表
            table = doc.add_table(word_row,word_column,style="Table Grid")
            # 写入表头
            table.cell(0,0).text = "序号"
            # 行
            for i in range(table_row):
                # 写入表头数据
                table.cell(i+1, 0).text = str(i + 1)
            # 列
            for i in range(table_column):
                table.cell(0, i+1).text = self.exp_data_3a.horizontalHeaderItem(i).text()
            # 写入内容
            for i in range(table_row):
                for j in range(table_column):
                    if(self.exp_data_3a.item(i,j) != None):
                        table.cell(i+1,j+1).text = self.exp_data_3a.item(i,j).text()

            # 表2
            doc.add_paragraph("")
            doc.add_paragraph("表2 粗糙管数据表")
            # 获取行列
            table_row_2 = self.exp_data_3b.rowCount()
            table_column_2 = self.exp_data_3b.columnCount()
            word_row_2 = self.exp_data_3b.rowCount() + 1
            word_column_2 = self.exp_data_3b.columnCount() + 1
            table_2 = doc.add_table(word_row_2, word_column_2, style="Table Grid")
            # 写入表头
            table_2.cell(0, 0).text = "序号"
            # 行
            for i in range(table_row_2):
                # 写入表头数据
                table_2.cell(i + 1, 0).text = str(i + 1)
            # 列
            for i in range(table_column_2):
                table_2.cell(0, i + 1).text = self.exp_data_3b.horizontalHeaderItem(i).text()
            # 写入内容
            for i in range(table_row_2):
                for j in range(table_column_2):
                    if (self.exp_data_3b.item(i, j) != None):
                        table_2.cell(i + 1, j + 1).text = self.exp_data_3b.item(i, j).text()

            # 表3
            doc.add_paragraph("")
            doc.add_paragraph("表3 阀门局部阻力系数测定数据表")
            # 获取行列
            table_row_3 = self.exp_data_3c.rowCount()
            table_column_3 = self.exp_data_3c.columnCount()
            word_row_3 = self.exp_data_3c.rowCount()
            word_column_3 = self.exp_data_3c.columnCount() + 1
            table_3 = doc.add_table(word_row_3, word_column_3, style="Table Grid")
            # 写入表头
            # 行
            for i in range(table_row_3):
                # 写入表头数据
                table_3.cell(i , 0).text = self.exp_data_3c.verticalHeaderItem(i).text()
            # 写入内容
            for i in range(table_row_3):
                for j in range(table_column_3):
                    if (self.exp_data_3c.item(i, j) != None):
                        table_3.cell(i , j + 1).text = self.exp_data_3c.item(i, j).text()

            # 获取保存路径
            file_path = pre_path + '/' + '实验三 流体阻力的测定.docx'
            doc.save(file_path)

            # 创建一个问答框
            self.box = QMessageBox(QMessageBox.Question, '提示', 'docx文件已保存，是否转换成pdf文件？')

            # 添加按钮
            yes = self.box.addButton('确定', QMessageBox.YesRole)
            no = self.box.addButton('取消', QMessageBox.NoRole)

            # 显示该问答框
            self.box.exec_()

            if self.box.clickedButton() == yes:
                print(1)
                # docx转pdf
                docx2pdf(file_path)
                QMessageBox.information(self, '提示', '成功生成pdf文件！', QMessageBox.Ok)
            else:
                print(2)
                pass

    # except:
    #     QMessageBox.critical(self, '错误', '导出失败', QMessageBox.Ok)