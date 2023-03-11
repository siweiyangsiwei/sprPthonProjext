from PyQt5.QtWidgets import QFileDialog, QMessageBox
from docx import Document


def import_data(self):
    file_path = QFileDialog.getExistingDirectory()
    return file_path

def write_9_docx(self,data, date, bgtemp, fgtemp, flow, fstemp, pressure, size, weight, area):
    # 完成实验报告
    document = Document('./resources/report/实验九 洞道干燥实验.docx')

    document.add_paragraph('实验日期' + str(date) + "\t干燥介质:热空气\t物料尺寸" + str(size) +
                           "\t物料的绝干质量" + str(weight) +
                           "\t干燥室截面积" + str(area) + "\t室前干球温度" + str(fgtemp) +
                           "\t室前湿球温度" + str(fstemp) + " \t室后干球温度" + str(bgtemp) +
                           "\t孔板流量计空气流量" + str(flow) + "\t风机出口压力" + str(pressure)
                           )

    document.add_paragraph("                表1 数据记录表")
    table1 = document.add_table(len(data[0]), len(data), style='Table Grid')
    heading_cells = table1.rows[0].cells
    heading_cells[0].text = "湿物料质量Gi(g)"
    heading_cells[1].text = "湿物料含水量Xi(kg水/kg绝干料)"
    heading_cells[2].text = "湿物料平均含水量(kg水/kg绝干料)"
    heading_cells[3].text = "汽化水分量△W(g)"
    heading_cells[4].text = "时间间隔△t(m,s)"
    heading_cells[5].text = "干燥速率U(kg/m²,s"

    for i in range(len(data)):
        for j in range(len(data[0]) - 1):
            if i > 1 and j > len(data[0]) - 3:
                continue
            table1.rows[j + 1].cells[i].text = str(data[i][j])

    document.add_picture('./data/img/exp_9_data_1.png')
    document.add_picture('./data/img/exp_9_data_2.png')

    # 保存文件
    pre_path = import_data(self)
    file_path = pre_path + '/' + '实验九 洞道干燥实验.docx'
    document.save(file_path)
    QMessageBox.information(self, '提示', '导出成功', QMessageBox.Ok)
