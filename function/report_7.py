from PyQt5.QtWidgets import QFileDialog, QMessageBox
from docx import Document


def import_data(self):
    file_path = QFileDialog.getExistingDirectory()
    return file_path

def write_7_docx(self, data1, data2, data3, HETP, R1, R2, R3, q1, q2, q3, date, num_of_pic, num_of_shai_ban,
             height_of_tian_liao,tower_type):
    # 实验报告
    document = Document('./resources/report/实验七 精馏实验.docx')
    document.add_paragraph('实验日期' + str(date) + "\t板式塔实际塔板数" + str(num_of_shai_ban) +
                           "\t填料塔填料层高度" + str(height_of_tian_liao) +
                           "\t图解法计算所得理论板数" + str(num_of_pic))
    document.add_paragraph("                表1 数据记录表")
    table1 = document.add_table(len(data1) + 1, 4, style='Table Grid')
    table1.rows[0].cells[1].text = '1'
    table1.rows[0].cells[2].text = '2'
    table1.rows[0].cells[3].text = '3'
    if not tower_type:
        for i in range(len(data1)):
            table1.rows[i + 1].cells[0].text = self.data_processing_7_table_1.verticalHeaderItem(i).data(0)

    else:
        for i in range(len(data1)):
            table1.rows[i + 1].cells[0].text = self.data_processing_7_table_2.verticalHeaderItem(i).data(0)

    for j in range(len(data1)):
        table1.rows[j + 1].cells[1].text = str(data1[j])
        table1.rows[j + 1].cells[2].text = str(data2[j])
        table1.rows[j + 1].cells[3].text = str(data3[j])
    document.add_paragraph("\n")

    document.add_paragraph("                表2 计算结果")
    table2 = document.add_table(4, 4, style='Table Grid')
    table2.rows[0].cells[1].text = '1'
    table2.rows[0].cells[2].text = '2'
    table2.rows[0].cells[3].text = '3'
    for i in range(3):
        table2.rows[i + 1].cells[0].text = self.data_processing_7_table_3.verticalHeaderItem(i).data(0)
        table2.rows[1].cells[1].text = str(HETP)
        table2.rows[1].cells[2].text = str(HETP)
        table2.rows[1].cells[3].text = str(HETP)
        table2.rows[2].cells[1].text = str(R1)
        table2.rows[2].cells[2].text = str(R2)
        table2.rows[2].cells[3].text = str(R3)
        table2.rows[3].cells[1].text = str(q1)
        table2.rows[3].cells[2].text = str(q2)
        table2.rows[3].cells[3].text = str(q3)

    # 保存文件
    pre_path = import_data(self)
    file_path = pre_path + '/' + '实验七 精馏实验.docx'
    document.save(file_path)
    QMessageBox.information(self, '提示', '导出成功', QMessageBox.Ok)